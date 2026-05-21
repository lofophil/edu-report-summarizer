# -*- coding: utf-8 -*-
"""
02_docx_to_sections.py（5 模块版）

功能：
- 从 “报告_docx” 目录读取所有 .docx 文件；
- 只提取并归类为 5 个规范模块：
    研究问题
    核心概念
    研究目标和内容
    研究成果
    研究效果
- 每篇报告输出为一条 JSON line，写入：
    data/reports_sections.jsonl

说明：
- 原文中的“背景、研究综述、研究意义、研究方法、研究措施与活动”等部分
  不再单独作为模块输出，但全文内容仍会进入 full_text，用于后续训练。
"""

import re
import json
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from tqdm import tqdm


# =========================
# 路径配置
# =========================
BASE_DIR = Path(__file__).resolve().parent
REPORT_DOCX_DIR = BASE_DIR / "data"
DATA_DIR = BASE_DIR / "data"
DATA_DIR.mkdir(parents=True, exist_ok=True)
OUT_JSONL = DATA_DIR / "reports_sections.jsonl"


# =========================
# 规范模块名（固定 5 个）
# =========================
CANONICAL_SECTIONS = [
    "研究问题",
    "核心概念",
    "研究目标和内容",
    "研究成果",
    "研究效果",
]


# =========================
# 模块标题别名（可以根据实际再补充）
# =========================
SECTION_ALIASES = {
    "研究问题": [
        "研究问题", "核心问题", "主要问题", "关键问题",
        "研究的主要问题", "问题的提出", "问题提出",
        "存在问题", "问题分析", "问题与对策"
    ],
    "核心概念": [
        "核心概念", "概念界定", "相关概念界定", "基本概念界定",
        "主要概念界定", "术语界定", "名词界定", "核心概念界定",
        "相关概念", "核心术语界定"
    ],
    "研究目标和内容": [
        "研究目标和内容", "研究目标与内容", "研究目标及内容",
        "研究目的与内容", "研究目标", "研究目的",
        "研究内容", "主要内容", "研究的主要内容",
        "研究目标与主要内容", "研究目标与任务",
        "研究任务", "研究思路与实施内容",
        "研究过程与内容", "实施内容", "研究计划与内容"
    ],
    "研究成果": [
        "研究成果", "主要成果", "阶段成果", "课题成果",
        "课题研究成果", "研究收获", "研究结论",
        "主要结论", "研究进展与成果",
        "创新点", "特色与创新", "成果概述", "研究产出"
    ],
    "研究效果": [
        "研究效果", "实施效果", "应用效果", "实践效果", "推广效果",
        "应用价值", "成果应用", "成果推广", "社会效益",
        "教学效果", "实施成效", "研究成效",
        "预期目标达成情况", "课题实施效果", "综合效果"
    ],
}


# =========================
# 工具函数
# =========================
def normalize_heading(text: str) -> str:
    """
    对段落标题做规范化处理，便于匹配：
    - 去掉全角/半角空格
    - 去掉前面的序号（如“一、”“（一）”“1.” 等）
    - 去掉结尾的冒号、分号、句号等
    """
    if not text:
        return ""

    t = text.strip()
    # 去掉所有空白（含中文空格）
    t = re.sub(r"\s+", "", t)

    # 去掉开头序号：如 “一、”“(一)”“1.”“1）”“第1部分”等
    t = re.sub(
        r'^(第?[零一二三四五六七八九十百千\d]+[章节部分]?)'
        r'[、\.\．\)\）\-\s]*',
        '',
        t
    )

    # 去掉末尾的冒号、分号、句号等
    t = re.sub(r'[：:；;，,。.\s]+$', '', t)

    return t.strip()


def match_section(heading_text: str) -> Optional[str]:
    """
    把一个“看起来像标题”的文本，映射到 5 个规范模块之一。
    """
    if not heading_text:
        return None

    h = normalize_heading(heading_text)
    if not h:
        return None

    # 1) 完全一致优先
    for canon in CANONICAL_SECTIONS:
        if h == canon:
            return canon

    # 2) 别名匹配（只在 5 个模块内搜索）
    for canon, aliases in SECTION_ALIASES.items():
        for alias in aliases:
            alias_clean = re.sub(r'\s+', '', alias)
            if alias_clean and alias_clean in h:
                return canon

    return None


def looks_like_heading(paragraph) -> bool:
    """
    结合样式 & 文本长度粗略判断是否“像一个标题”。
    - 若 style 名称包含 'Heading' 或 '标题'，优先视为标题；
    - 文本太短（<= 25 字符）也倾向看作标题；
    """
    text = paragraph.text.strip()
    if not text:
        return False

    try:
        style_name = paragraph.style.name or ""
    except Exception:
        style_name = ""

    style_name = str(style_name)

    # 样式判断
    if ("Heading" in style_name) or ("标题" in style_name):
        return True

    # 纯长度判断：太长的一般不是一级标题
    if len(text) <= 25:
        return True

    return False


# =========================
# 核心解析函数
# =========================
def docx_to_sections(path: Path) -> Dict:
    """
    解析单个 docx：
    返回结构：
    {
      "doc_id": "...",
      "title": "...",
      "sections": {
          "研究问题": "...",
          "核心概念": "...",
          "研究目标和内容": "...",
          "研究成果": "...",
          "研究效果": "..."
      },
      "full_text": "全文拼接"
    }
    """
    doc = Document(str(path))

    # 初始化：5 个模块
    sections_buf: Dict[str, List[str]] = {sec: [] for sec in CANONICAL_SECTIONS}
    full_text_parts: List[str] = []

    # 粗略提取“报告标题”：取第一个不太长的非空段落
    title = ""
    for p in doc.paragraphs:
        t = p.text.strip()
        if t and len(t) <= 40:  # 经验阈值
            title = t
            break

    current_section: Optional[str] = None

    for para in doc.paragraphs:
        raw = para.text
        text = raw.strip()
        if not text:
            continue

        # full_text：所有非空段落都拼进去
        full_text_parts.append(text)

        # 判断是不是“标题段落”
        if looks_like_heading(para):
            matched = match_section(text)
            if matched is not None:
                current_section = matched
                continue  # 标题本身不写入内容
            else:
                # 看起来像标题，但没匹配到 5 个模块：视为普通正文（只进 full_text）
                continue

        # 若不是标题，且已经有当前模块，就归入该模块
        if current_section is not None:
            sections_buf[current_section].append(text)
        # 如果 current_section 仍为 None，则该段只进入 full_text，不进入任何模块

    # 把每个模块的列表拼成长文本
    sections_text = {
        sec: "\n".join(paras).strip()
        for sec, paras in sections_buf.items()
    }

    full_text = "\n".join(full_text_parts).strip()

    record = {
        "doc_id": path.stem,
        "title": title,
        "sections": sections_text,
        "full_text": full_text,
    }
    return record


# =========================
# 主程序入口
# =========================
def main():
    if not REPORT_DOCX_DIR.exists():
        print(f"[ERROR] 报告目录不存在：{REPORT_DOCX_DIR}")
        return

    docx_files = sorted(REPORT_DOCX_DIR.glob("*.docx"))
    if not docx_files:
        print(f"[WARN] 在 {REPORT_DOCX_DIR} 下没有找到任何 .docx 文件。")
        return

    print(f"[INFO] 共找到 {len(docx_files)} 个 docx 文件。")
    print(f"[INFO] 输出 JSONL：{OUT_JSONL}")

    n_ok, n_err = 0, 0
    with OUT_JSONL.open("w", encoding="utf-8") as fw:
        for p in tqdm(docx_files, desc="解析 docx -> sections"):
            try:
                rec = docx_to_sections(p)
                fw.write(json.dumps(rec, ensure_ascii=False) + "\n")
                n_ok += 1
            except Exception as e:
                print(f"[ERROR] 解析失败: {p} -> {e}")
                n_err += 1

    print(f"[DONE] 成功解析 {n_ok} 篇，失败 {n_err} 篇。输出文件：{OUT_JSONL}")


if __name__ == "__main__":
    main()
